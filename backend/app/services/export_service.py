from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from sqlalchemy.orm import Session
from app.models.builder_session import BuilderSession


class ExportService:

    def __init__(self, db: Session):
        self.db = db

    def export_pdf_for_session(self, session: BuilderSession) -> bytes:
        buffer = BytesIO()
        resume_data = session.resume_data or {}
        doc = SimpleDocTemplate(
            buffer, pagesize=A4,
            rightMargin=20*mm, leftMargin=20*mm,
            topMargin=20*mm, bottomMargin=20*mm
        )
        styles = getSampleStyleSheet()
        story = []

        name_style = ParagraphStyle('Name', fontSize=20, fontName='Helvetica-Bold', spaceAfter=4)
        story.append(Paragraph(resume_data.get('name', 'Your Name'), name_style))

        contact_parts = [str(resume_data[f]) for f in ['email', 'phone', 'linkedin', 'portfolio'] if resume_data.get(f)]
        if contact_parts:
            contact_style = ParagraphStyle('Contact', fontSize=9, textColor=colors.grey, spaceAfter=8)
            story.append(Paragraph(' | '.join(contact_parts), contact_style))

        story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
        story.append(Spacer(1, 6))

        section_style = ParagraphStyle('Section', fontSize=12, fontName='Helvetica-Bold', spaceAfter=4, spaceBefore=10)
        body_style = ParagraphStyle('Body', fontSize=10, spaceAfter=3, leading=14)
        bullet_style = ParagraphStyle('Bullet', fontSize=10, spaceAfter=2, leftIndent=12, leading=14)

        def add_section(title, items, bullet=False):
            if not items:
                return
            story.append(Paragraph(title.upper(), section_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
            story.append(Spacer(1, 4))
            for item in (items if isinstance(items, list) else [items]):
                style = bullet_style if bullet else body_style
                prefix = '- ' if bullet else ''
                story.append(Paragraph(f'{prefix}{item}', style))

        skills = resume_data.get('skills', [])
        if skills:
            story.append(Paragraph('SKILLS', section_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
            story.append(Spacer(1, 4))
            story.append(Paragraph(' | '.join(skills) if isinstance(skills, list) else str(skills), body_style))

        add_section('Education', resume_data.get('education', []))
        add_section('Experience', resume_data.get('experience', []), bullet=True)
        add_section('Projects', resume_data.get('projects', []), bullet=True)
        add_section('Certifications', resume_data.get('certifications', []), bullet=True)
        add_section('Achievements', resume_data.get('achievements', []), bullet=True)

        doc.build(story)
        return buffer.getvalue()

    def export_docx_for_session(self, session: BuilderSession) -> bytes:
        resume_data = session.resume_data or {}
        buffer = BytesIO()
        doc = Document()

        doc.add_heading(resume_data.get('name', 'Your Name'), level=1)

        contact_parts = [str(resume_data[f]) for f in ['email', 'phone', 'linkedin', 'portfolio'] if resume_data.get(f)]
        if contact_parts:
            doc.add_paragraph(' | '.join(contact_parts))

        def add_section(title, items, bullet=False):
            if not items:
                return
            doc.add_heading(title, level=2)
            for item in (items if isinstance(items, list) else [items]):
                if bullet:
                    doc.add_paragraph(str(item), style='List Bullet')
                else:
                    doc.add_paragraph(str(item))

        skills = resume_data.get('skills', [])
        if skills:
            doc.add_heading('Skills', level=2)
            doc.add_paragraph(', '.join(skills) if isinstance(skills, list) else str(skills))

        add_section('Education', resume_data.get('education', []))
        add_section('Experience', resume_data.get('experience', []), bullet=True)
        add_section('Projects', resume_data.get('projects', []), bullet=True)
        add_section('Certifications', resume_data.get('certifications', []), bullet=True)
        add_section('Achievements', resume_data.get('achievements', []), bullet=True)

        doc.save(buffer)
        return buffer.getvalue()